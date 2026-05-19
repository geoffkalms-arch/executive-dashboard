"""Generate bottom-load fuel skid executive briefing note as a Word .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1b, 0x2a, 0x4a)
MID    = RGBColor(0x2e, 0x41, 0x72)
RED    = RGBColor(0xc0, 0x39, 0x2b)
AMBER  = RGBColor(0xd6, 0x89, 0x10)
GREEN  = RGBColor(0x1a, 0x7a, 0x3f)
WHITE  = RGBColor(0xff, 0xff, 0xff)
MUTED  = RGBColor(0x5a, 0x62, 0x78)
BG_RED = RGBColor(0xfd, 0xe8, 0xe8)
BG_GRN = RGBColor(0xd4, 0xed, 0xda)
BG_NVY = RGBColor(0xea, 0xf0, 0xfb)
BG_AMB = RGBColor(0xfe, 0xf3, 0xcd)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def add_rule(doc, color=MUTED):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_heading(doc, text, color=NAVY, size=14, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    add_rule(doc, color)
    return p


def add_body(doc, text, size=10.5, space_before=2, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def add_callout(doc, label, text, bg, label_color):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label.upper())
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.color.rgb = label_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def badge(text, color):
    return (text, True, color)


def add_table(doc, headers, rows, header_bg=NAVY, alt_bg=BG_NVY):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 1 else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if isinstance(cell_text, tuple):
                txt, bold, color = cell_text
                r = p.add_run(txt)
                r.bold = bold
                r.font.size = Pt(9)
                if color:
                    r.font.color.rgb = color
            else:
                r = p.add_run(str(cell_text))
                r.font.size = Pt(9)

    doc.add_paragraph()
    return tbl


# ═══════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21.0)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ── Header band ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
r = p.add_run('[Your Company Name]  |  Fuel Infrastructure Solutions')
r.font.size = Pt(9)
r.font.color.rgb = MUTED

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('EXECUTIVE BRIEFING NOTE')
r.bold = True
r.font.size = Pt(8)
r.font.color.rgb = RED

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Bottom Load Fuel Skid Upgrade — Road Tanker Loading Safety')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = NAVY

# meta line
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
for label, value in [('For:', '[Customer Name]'), ('By:', '[Your Company Name]'), ('Date:', 'May 2026'), ('Ref:', 'BLS-2026-001-EB')]:
    r1 = p.add_run(f'  {label}  ')
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    r2 = p.add_run(value)
    r2.bold = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = NAVY

add_rule(doc, NAVY)

# ── The issue ──
add_heading(doc, 'The Issue in One Paragraph', size=13, space_before=10)

add_body(doc,
    "[Customer Name]'s current road tanker loading facility uses top-fill loading for diesel and "
    "unleaded petrol. Every fill operation simultaneously creates four uncontrolled safety hazards: "
    "the risk of tanker overfill, the release of flammable and toxic petrol vapours directly to "
    "atmosphere, the risk of electrostatic sparks igniting those vapours, and the requirement for "
    "drivers to work at 2.7–3.5 m height every time they load. All four hazards are rated EXTREME "
    "under the Australian Standard risk matrix. The facility is non-compliant with AS 1940:2017, the "
    "Work Health and Safety Act 2011, and at risk under state dangerous goods and environmental "
    "legislation. The tankers are already set up for bottom fill. The solution is available, proven, "
    "and straightforward."
)

# ── 4 hazards table ──
add_heading(doc, 'The Four Hazards — At a Glance', size=13, space_before=10)

add_table(doc,
    ['Hazard', 'Current Risk', 'After Upgrade'],
    [
        (('Overfill', True, NAVY),
         badge('EXTREME — operator vigilance only, no auto shut-off', RED),
         badge('LOW — automatic OPV + electronic interlock', GREEN)),
        (('Vapour release', True, NAVY),
         badge('EXTREME — open hatch, continuous venting every fill', RED),
         badge('LOW — closed system, 95%+ vapour captured & recovered', GREEN)),
        (('Electrostatic spark', True, NAVY),
         badge('EXTREME — manual bonding cable, no interlock, often skipped', RED),
         badge('LOW — coded electronic grounding verification with pump interlock', GREEN)),
        (('Working at height', True, NAVY),
         badge('EXTREME — 2.7–3.5 m access required every fill, every day', RED),
         badge('LOW — all operations at ground level, fall hazard eliminated', GREEN)),
    ]
)

# ── What goes in ──
add_heading(doc, 'What Goes In', size=13, space_before=10)

add_body(doc,
    'A Bottom Load Skid is a ground-level loading station that connects to the tanker from below '
    'via a sealed, dry-break API coupler. The tanker remains completely closed throughout the fill. '
    'The system includes:'
)

components = [
    ('API dry-break bottom-fill coupler', '— sealed connection, no spill on disconnect'),
    ('Automatic overfill prevention (OPV)', '— product stops flowing when compartment is full, automatically'),
    ('Vapour recovery system (VRS)', '— displaced vapours return to the storage tank as liquid, not to air'),
    ('Electronic grounding verification', '— product cannot flow unless tanker is confirmed grounded; monitored continuously'),
    ('Preset metered delivery', '— accurate volume control and stock reconciliation'),
    ('Ground-level emergency stop', '— immediate, accessible, at hand'),
]

for comp, desc in components:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(comp)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(f'  {desc}')
    r2.font.size = Pt(10)

doc.add_paragraph()

# ── Regulatory ──
add_heading(doc, 'Regulatory Position', size=13, space_before=10)

add_body(doc, 'The current system is non-compliant with:')

regs = [
    ('AS 1940:2017 Cl. 5.5', '— requires a grounding interlock for tanker loading; a manual cable without an interlock does not satisfy this requirement'),
    ('WHS Act 2011 s.19', '— requires elimination of foreseeable risks where reasonably practicable'),
    ('WHS Regulations 2017 Reg. 78', '— working at height controls (handrails alone may not satisfy this where operators must lean over to access the hatch)'),
]

for reg, desc in regs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(reg)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RED
    r2 = p.add_run(f'  {desc}')
    r2.font.size = Pt(10)

doc.add_paragraph()

add_callout(doc,
    'Insurance Risk',
    'Many industrial and public liability insurers now require evidence of AS 1940 compliance as a '
    'condition of cover. Facilities found non-compliant at the time of an incident may face claims '
    'being declined or coverage voidance.',
    BG_AMB, AMBER
)

# ── Cost ──
add_heading(doc, 'Cost Perspective', size=13, space_before=10)

add_table(doc,
    ['Item', 'Range (AUD)'],
    [
        ('Bottom-load skid system — 2 to 4 bays, installed', '$180,000 – $450,000'),
        (('Single WHS Category 1 prosecution penalty', True, RED), ('Up to $3,000,000', True, RED)),
        (('Tanker fire / facility loss', True, RED), ('$500,000 – $10,000,000+', True, RED)),
        (('Environmental prosecution + remediation', True, RED), ('$25,000 – $1,000,000+', True, RED)),
        ('Ongoing product loss — vapour & overfill (per year)', '$5,000 – $40,000 per year'),
    ]
)

add_body(doc,
    'The upgrade cost is a one-time, planned capital investment. The alternative is a much larger, '
    'unplanned, and potentially uninsured liability.'
)

# ── Next steps ──
add_heading(doc, 'What Happens Next', size=13, space_before=10)

steps = [
    ('Authorise a site survey', '[Your Company Name] attends site to confirm scope — typically 1 day'),
    ('Detailed design and fixed-price quotation', 'prepared within 2–3 weeks of survey'),
    ('Equipment ordered', 'lead time typically 8–14 weeks for fabricated skids'),
    ('Installation', 'coordinated bay-by-bay to maintain throughput during works'),
    ('Commissioning, operator training, and handover documentation', 'included'),
]

for i, (step, detail) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{i}.  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(step)
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = NAVY
    r3 = p.add_run(f' — {detail}')
    r3.font.size = Pt(10.5)

doc.add_paragraph()

# ── Bottom line ──
add_callout(doc,
    'Bottom Line',
    'The risk is real, the liability is current, and the solution is available. '
    'The tankers are already equipped for bottom fill. The question is not whether to upgrade — it is when.\n\n'
    'Contact [Your Company Name] to arrange a site survey.',
    BG_GRN, GREEN
)

# ── Footer ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top_el = OxmlElement('w:top')
top_el.set(qn('w:val'), 'single')
top_el.set(qn('w:sz'), '4')
top_el.set(qn('w:space'), '1')
top_el.set(qn('w:color'), f"{MUTED[0]:02X}{MUTED[1]:02X}{MUTED[2]:02X}")
pBdr.append(top_el)
pPr.append(pBdr)
r = p.add_run('BLS-2026-001-EB  |  [Your Company Name]  |  May 2026  |  Commercial in Confidence')
r.font.size = Pt(8)
r.font.color.rgb = MUTED

out = '/home/user/executive-dashboard/Bottom-Load-Executive-Briefing-Note.docx'
doc.save(out)
print(f'Saved: {out}')
