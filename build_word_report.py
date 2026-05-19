"""Generate bottom-load fuel skid safety report as a Word .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ──────────────────────────────────────────────
NAVY   = RGBColor(0x1b, 0x2a, 0x4a)
MID    = RGBColor(0x2e, 0x41, 0x72)
RED    = RGBColor(0xc0, 0x39, 0x2b)
AMBER  = RGBColor(0xd6, 0x89, 0x10)
GREEN  = RGBColor(0x1a, 0x7a, 0x3f)
WHITE  = RGBColor(0xff, 0xff, 0xff)
LIGHT  = RGBColor(0xea, 0xf0, 0xfb)
MUTED  = RGBColor(0x5a, 0x62, 0x78)
BG_RED = RGBColor(0xfd, 0xe8, 0xe8)
BG_AMB = RGBColor(0xfe, 0xf3, 0xcd)
BG_GRN = RGBColor(0xd4, 0xed, 0xda)
BG_NVY = RGBColor(0xea, 0xf0, 0xfb)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_heading(doc, text, level=1, color=NAVY, space_before=18, space_after=8, size=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if level == 1:
        # underline via border on paragraph bottom
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), f"{NAVY[0]:02X}{NAVY[1]:02X}{NAVY[2]:02X}")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_body(doc, text, space_before=2, space_after=6, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = NAVY
        r = p.add_run(text)
        r.font.size = Pt(size)
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
    return p


def add_callout(doc, label, text, bg: RGBColor, label_color: RGBColor):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label.upper())
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.color.rgb = label_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()  # spacer


def add_risk_table(doc, headers, rows, header_bg=NAVY, alt_bg=LIGHT):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = alt_bg if ri % 2 == 1 else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if isinstance(cell_text, tuple):
                # (text, bold, color)
                txt, bold, color = cell_text
                r = p.add_run(txt)
                r.bold = bold
                r.font.size = Pt(9)
                if color:
                    r.font.color.rgb = color
            else:
                r = p.add_run(str(cell_text))
                r.font.size = Pt(9)

    doc.add_paragraph()
    return tbl


def badge(text, color):
    """Return a (text, bold, color) tuple for use in risk tables."""
    return (text, True, color)


# ═══════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── Default paragraph style ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
r = p.add_run('[Your Company Name]  |  Fuel Infrastructure Solutions')
r.font.size = Pt(9)
r.font.color.rgb = MUTED

p = doc.add_paragraph()
r = p.add_run('SAFETY ASSESSMENT & BUSINESS CASE')
r.bold = True
r.font.size = Pt(9)
r.font.color.rgb = RED
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
r = p.add_run('Bottom Load Fuel Skid\nUpgrade Program')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(10)

p = doc.add_paragraph()
r = p.add_run('Road Tanker Loading — Transition from Top Fill to Bottom Fill Loading Systems')
r.font.size = Pt(14)
r.font.color.rgb = MID
p.paragraph_format.space_after = Pt(40)

meta_lines = [
    ('Prepared for:', '[Customer Name]'),
    ('Prepared by:', '[Your Company Name]'),
    ('Date:', 'May 2026'),
    ('Document Ref:', 'BLS-2026-001 Rev A'),
    ('Classification:', 'Commercial in Confidence'),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{label}  ')
    r1.font.size = Pt(10)
    r1.font.color.rgb = MUTED
    r2 = p.add_run(value)
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = NAVY

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    'This document is prepared for the exclusive use of [Customer Name]. Contents are based on site '
    'information provided and applicable Australian Standards and dangerous goods legislation current '
    'at date of issue. Specific engineering and regulatory advice should be obtained prior to installation.'
)
r.font.size = Pt(8.5)
r.font.color.rgb = MUTED
p.paragraph_format.space_before = Pt(30)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
add_heading(doc, '1.  Executive Summary')

add_body(doc,
    '[Customer Name] currently operates a top-fill road tanker loading facility with 2–4 loading bays '
    'handling diesel and unleaded petrol at a medium-volume regional depot (10–30 fills per day). This '
    'assessment identifies four critical, unmitigated or inadequately mitigated hazard categories present '
    'in the current configuration and demonstrates that the installation of a modern API Bottom Load Skid '
    'system will substantially eliminate each hazard class.'
)
add_body(doc,
    'The recommendation is to proceed with a full bottom-load conversion, incorporating: API bottom-fill '
    'couplers, overfill prevention valves (OPV), a vapour recovery system (VRS), and an electronic '
    'grounding verification system. This will bring the facility into full compliance with applicable '
    'Australian Standards and dangerous goods legislation, materially reduce risk of fire, explosion, '
    'environmental contamination, and serious injury, and reduce long-term operating costs.'
)

add_callout(doc,
    'Key Finding',
    'Under the current top-fill configuration, all four primary hazard categories — overfill, vapour '
    'release, electrostatic ignition, and working at height — are simultaneously present during every '
    'tanker fill. The likelihood of a serious incident is rated High on the Australian Standard risk '
    'matrix. This represents an unacceptable residual risk under the Work Health and Safety Act 2011 '
    '(Cth) and applicable state dangerous goods legislation.',
    BG_RED, RED
)

# ════════════════════════════════════════════════════════════════
# 2. REGULATORY FRAMEWORK
# ════════════════════════════════════════════════════════════════
add_heading(doc, '2.  Regulatory & Standards Framework')

add_body(doc,
    'The following legislation and standards govern the storage, handling, and loading of Class 3 '
    'flammable liquids (diesel and unleaded petrol) at road tanker loading facilities in Australia. '
    'Compliance is mandatory, not optional, and enforcement authority rests with both the relevant '
    'state Work Health and Safety regulator and the Dangerous Goods regulator.'
)

regs = [
    ('Work Health and Safety Act 2011 (Cth) / Mirror Acts',
     'Section 19 — Primary duty of care. A person conducting a business or undertaking (PCBU) must '
     'ensure, so far as is reasonably practicable, the health and safety of workers and others. This '
     'is an absolute obligation to eliminate risks where possible, or where not possible, to minimise '
     'them to the lowest practicable level. Known, controllable hazards that are not mitigated expose '
     'the PCBU to prosecution.'),
    ('AS 1940:2017 — The Storage and Handling of Flammable and Combustible Liquids',
     'The principal Australian Standard for liquid fuel handling. Mandates requirements for overfill '
     'prevention, vapour control, earthing/bonding of transfer equipment, and tanker loading '
     'procedures. Top-fill systems must comply with specific controls; bottom-fill systems are '
     'recognised as the preferred method meeting these requirements inherently.'),
    ('Dangerous Goods Legislation — State/Territory',
     'All states require manifested tanker loading to comply with relevant codes. Loading sites must '
     'be designed to prevent uncontrolled release of dangerous goods and must provide safe working '
     'conditions for transport operators. Relevant acts include the Dangerous Substances Act (Vic), '
     'Dangerous Goods Safety Management Act (Qld), and equivalents in other states.'),
    ('Australian Dangerous Goods Code (ADG Code), 9th Edition',
     'Governs the road transport of dangerous goods. Tankers used for Class 3 liquids operating in '
     'bottom-fill configuration must be loaded at a compliant bottom-fill station. This is '
     'increasingly a legal and insurance requirement for transport operators.'),
    ('AS/NZS 60079 Series — Explosive Atmospheres',
     'Defines hazardous area zones around loading points. Zone 1 (area where explosive atmospheres '
     'are likely to occur) exists around open top-fill hatches. Equipment selection, earthing, and '
     'vapour control in these zones must comply with this series.'),
    ('Safe Work Australia — Code of Practice: Managing Risks of Hazardous Chemicals',
     'Requires a hierarchy of controls to manage chemical hazards. Elimination or substitution '
     '(i.e. replacing the hazardous top-fill method with an inherently safer bottom-fill method) is '
     'the highest-ranked control and must be implemented where reasonably practicable.'),
]

for title, body in regs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = MID
    add_body(doc, body, space_before=0, space_after=4, size=10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 3. CURRENT SYSTEM RISK ASSESSMENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '3.  Current System — Risk Assessment')
add_body(doc,
    'The following section details each identified hazard in the current top-fill loading configuration. '
    'Risk ratings are based on the AS/NZS ISO 31000 risk management framework using a 5×5 '
    'likelihood/consequence matrix.'
)

# ── 3.1 Overfill ──
add_heading(doc, '3.1  Hazard 1: Overfill of Road Tanker', level=2, color=MID, size=13, space_before=10)

add_heading(doc, 'How it occurs', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_body(doc,
    'In a top-fill loading configuration, product is gravity-fed or pumped into an open top hatch on '
    'the tanker. Flow control relies entirely on the operator manually monitoring the fill and closing '
    'a valve or activating a pump stop. There is no automatic overfill shut-off connected to the '
    'tanker\'s own level sensors.', size=10
)

add_heading(doc, 'Consequences', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_bullet(doc, 'Uncontrolled release of diesel or petrol directly onto the hardstand and into stormwater drainage, risking soil and groundwater contamination and triggering EPA notification obligations.', bold_prefix='Spill to ground and stormwater:  ')
add_bullet(doc, 'Petrol vapour from a spill at ambient temperature can rapidly form an explosive atmosphere across a wide area. A single ignition source — a running vehicle engine, a static spark, a cigarette — can detonate the mixture.', bold_prefix='Fire and explosion:  ')
add_bullet(doc, 'Even partial overfills represent direct financial loss plus remediation and clean-up costs.', bold_prefix='Product loss:  ')
add_bullet(doc, 'An uncontrolled release of Class 3 flammable liquids to stormwater is a notifiable environmental incident in all Australian states.', bold_prefix='Regulatory breach:  ')

add_callout(doc,
    'Specific Risk — Petrol (ULP)',
    'Unleaded petrol has a flash point of −43°C, meaning it produces ignitable vapours at all ambient '
    'temperatures encountered in Australian conditions. An overfill of even 20–30 litres can produce '
    'an explosive vapour cloud extending many metres from the spill point in calm conditions.',
    BG_RED, RED
)

add_heading(doc, 'Current controls', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_bullet(doc, 'Operator visual monitoring — highly fallible; distraction, fatigue, and obscured sight lines are common.')
add_bullet(doc, 'Tanker-mounted ullage sight glasses — may not be visible from the fill platform.')
add_bullet(doc, 'No automatic or electronic overfill prevention integrated into the loading system.')

add_heading(doc, 'Risk Rating — Current State', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_risk_table(doc,
    ['Parameter', 'Rating', 'Basis'],
    [
        ('Likelihood', badge('Likely (4)', RED), 'Operator-dependent control; human error occurs regularly across the industry'),
        ('Consequence', badge('Major (4)', RED), 'Fire/explosion potential; environmental prosecution; serious injury or fatality'),
        (('Residual Risk', True, NAVY), badge('EXTREME (16)', RED), 'Immediate action required'),
    ]
)

# ── 3.2 Vapour ──
add_heading(doc, '3.2  Hazard 2: Uncontrolled Vapour Release', level=2, color=MID, size=13, space_before=10)

add_heading(doc, 'How it occurs', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_body(doc,
    'Top-fill loading requires the tanker hatch to be open throughout the fill process. As product '
    'enters the compartment, the vapour space above the liquid is displaced and vented directly to '
    'atmosphere through the open hatch. At a 2–4 bay facility loading 10–30 tankers per day, this '
    'is a continuous, high-volume vapour source during operating hours.', size=10
)

add_heading(doc, 'Consequences', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_bullet(doc, 'Petrol vapours discharged from an open hatch create an uncontrolled hazardous zone (Zone 1 per AS/NZS 60079) immediately above and around the fill point. Ignition sources anywhere in this zone can cause ignition.', bold_prefix='Fire and explosion risk:  ')
add_bullet(doc, 'Prolonged exposure to benzene (present in ULP) and other volatile hydrocarbons is a known carcinogen risk. Australian WHS regulations require atmospheric monitoring and control of occupational exposure limits (OELs).', bold_prefix='Health hazard to operators:  ')
add_bullet(doc, 'Australian National Pollutant Inventory (NPI) requires reporting of volatile organic compound (VOC) emissions above threshold quantities. Uncontrolled vapour venting from top-fill operations at a medium-volume site is likely to exceed NPI reporting thresholds.', bold_prefix='Environmental — VOC emissions:  ')
add_bullet(doc, 'Vented vapours represent direct, unrecovered product loss.', bold_prefix='Product loss:  ')

add_callout(doc,
    'Industry Data',
    'Studies by the US EPA and European Environment Agency have found that top-fill loading of petroleum '
    'products emits approximately 1.0–1.4 kg of VOCs per cubic metre of product transferred. At '
    'medium-volume throughput, this represents hundreds of kilograms of hydrocarbon vapour released per year.',
    BG_AMB, AMBER
)

add_heading(doc, 'Risk Rating — Current State', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_risk_table(doc,
    ['Parameter', 'Rating', 'Basis'],
    [
        ('Likelihood', badge('Almost Certain (5)', RED), 'Vapour release is inherent and unavoidable with every top-fill operation'),
        ('Consequence', badge('Major (4)', RED), 'Ignition risk; worker health impact; regulatory non-compliance'),
        (('Residual Risk', True, NAVY), badge('EXTREME (20)', RED), 'Immediate action required'),
    ]
)

# ── 3.3 Grounding ──
add_heading(doc, '3.3  Hazard 3: Electrostatic Discharge / Spark Ignition', level=2, color=MID, size=13, space_before=10)

add_heading(doc, 'How it occurs', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_body(doc,
    'Road tankers accumulate significant electrostatic charge during travel and during product movement '
    'through pipework. When an ungrounded tanker is approached with a fill hose or when the fill stream '
    'contacts the product in the compartment, electrostatic discharge (ESD) can occur. In a Zone 1 or '
    'Zone 2 hazardous atmosphere — which top-fill loading creates by definition — a single ESD event '
    'can be sufficient to ignite the vapour-air mixture.', size=10
)
add_body(doc,
    'This is not a theoretical risk. Multiple tanker fires and explosions worldwide have been attributed '
    'to ESD during top-fill loading, including several fatalities in the Australian petroleum '
    'distribution industry over recent decades.', size=10
)

add_heading(doc, 'Current controls and their limitations', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_bullet(doc, 'Operators are expected to clip an earthing cable to the tanker before opening hatches. This is a procedural control only — it is regularly skipped due to time pressure, lack of familiarity, or worn/damaged equipment. There is no verification that bonding has occurred before product can flow.', bold_prefix='Manual bonding cable:  ')
add_bullet(doc, 'In the current system, product can flow regardless of whether the tanker is bonded or grounded. There is no automatic prevention of product flow if grounding is absent.', bold_prefix='No interlock:  ')
add_bullet(doc, 'There is no system to detect dangerous charge levels on the tanker or to alert the operator.', bold_prefix='No ESD detection:  ')

add_callout(doc,
    'Critical Gap',
    'AS 1940:2017 Clause 5.5 requires that road tanker loading systems prevent product transfer unless '
    'effective earthing/bonding of the tanker is confirmed. A manual cable without an interlock does '
    'not satisfy this requirement. This is a current, auditable non-compliance.',
    BG_RED, RED
)

add_heading(doc, 'Risk Rating — Current State', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_risk_table(doc,
    ['Parameter', 'Rating', 'Basis'],
    [
        ('Likelihood', badge('Likely (4)', RED), 'Procedural reliance only; bonding skipped or inadequate regularly across industry'),
        ('Consequence', badge('Catastrophic (5)', RED), 'Ignition in Zone 1 atmosphere = explosion; multiple fatality potential'),
        (('Residual Risk', True, NAVY), badge('EXTREME (20)', RED), 'Highest priority action required'),
    ]
)

# ── 3.4 Working at height ──
add_heading(doc, '3.4  Hazard 4: Working at Height', level=2, color=MID, size=13, space_before=10)

add_heading(doc, 'How it occurs', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_body(doc,
    'Top-fill loading requires operators (typically the tanker driver) to climb to the top of the tanker '
    '— typically 2.7 m to 3.5 m above ground level — to open hatches, insert the fill pipe, and monitor '
    'the fill level. This is required for every fill, every day, in all weather conditions.', size=10
)

add_heading(doc, 'Consequences', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_bullet(doc, 'Falls from tanker tops are a documented cause of serious injury and fatality in the fuel distribution industry. Tank surfaces are often contaminated with product residue, making them slippery. Operations frequently occur early morning, in wet conditions, or after long drives.', bold_prefix='Falls from height:  ')
add_bullet(doc, 'A fall that results in the operator losing hold of the fill pipe can simultaneously create a product spill and a victim unable to evacuate — dramatically worsening the consequence of a concurrent fire or spill incident.', bold_prefix='Secondary ignition scenario:  ')
add_bullet(doc, 'WHS Regulations 2017, Regulation 78, requires that any work at height where a fall of 2 m or more could occur must be managed with a fall prevention device. Handrails alone may not satisfy this requirement if the operator must lean over the rail to perform the task.', bold_prefix='WHS obligation:  ')

add_callout(doc,
    'Note on Existing Handrails',
    'Handrails on a loading gantry reduce but do not eliminate the fall risk. If operators must reach '
    'over, around, or through handrails to access the tanker hatch — which is common on non-standard '
    'tanker configurations — the handrail provides no protection at the point of highest risk. The '
    'correct control is to eliminate the requirement to work at height entirely.',
    BG_AMB, AMBER
)

add_heading(doc, 'Risk Rating — Current State', level=3, color=NAVY, size=11, space_before=6, space_after=4)
add_risk_table(doc,
    ['Parameter', 'Rating', 'Basis'],
    [
        ('Likelihood', badge('Likely (4)', RED), 'Every fill requires height access; fall probability increases with volume and fatigue'),
        ('Consequence', badge('Major (4)', RED), 'Serious injury or fatality from fall; compounding fire/spill risk'),
        (('Residual Risk', True, NAVY), badge('EXTREME (16)', RED), 'Immediate action required'),
    ]
)

# ════════════════════════════════════════════════════════════════
# 4. RISK SUMMARY
# ════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Risk Summary — Current vs Proposed')

add_risk_table(doc,
    ['Hazard', 'Current Risk Rating', 'Post-Upgrade Risk Rating', 'Risk Reduction'],
    [
        ('1. Tanker Overfill',          badge('EXTREME (16)', RED), badge('LOW (2)', GREEN), 'Automatic OPV + API interlock eliminates operator reliance'),
        ('2. Vapour Release',           badge('EXTREME (20)', RED), badge('LOW (3)', GREEN), 'Closed-loop VRS captures >95% of displaced vapours'),
        ('3. Electrostatic / Grounding',badge('EXTREME (20)', RED), badge('LOW (2)', GREEN), 'Electronic grounding verification with pump interlock'),
        ('4. Working at Height',        badge('EXTREME (16)', RED), badge('LOW (1)', GREEN), 'All loading at ground level — fall hazard eliminated'),
    ]
)

add_callout(doc,
    'Outcome',
    'The proposed bottom-load system reduces all four hazard categories from EXTREME to LOW. This is '
    'consistent with the hierarchy of controls under WHS legislation: hazards are eliminated or '
    'engineered out, not managed through procedures and personal protective equipment.',
    BG_GRN, GREEN
)

# ════════════════════════════════════════════════════════════════
# 5. PROPOSED SYSTEM
# ════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Proposed Bottom Load System — Description & Benefits')

add_heading(doc, '5.1  How Bottom Loading Works', level=2, color=MID, size=13, space_before=10)
add_body(doc,
    'A bottom load system delivers fuel to the tanker through a dedicated API bottom-fill inlet located '
    'on the underside of each tanker compartment. The tanker remains closed throughout the fill. The '
    'loading skid connects to the tanker via a dry-break API coupler, which forms a sealed, drip-free '
    'connection. Product flows under pressure from ground level into the tanker from below, displacing '
    'vapour back through a closed vapour recovery line. All operations are performed at ground level by '
    'the driver.', size=10
)

add_heading(doc, '5.2  Key System Components', level=2, color=MID, size=13, space_before=10)

components = [
    ('API Dry-Break Bottom Fill Coupler',
     'The API coupler is a standardised, self-sealing quick-connect fitting that mates with the '
     'tanker\'s bottom-fill inlet. When disconnected, both the coupler and the tanker inlet valve '
     'close simultaneously, preventing any spillage. In a bottom-fill system, disconnection under '
     'flow cannot cause a spill — fundamentally different from top fill.'),
    ('Overfill Prevention Valve (OPV) / Overfill Protection System (OPS)',
     'Each tanker compartment fitted for bottom-fill has an Overfill Prevention Valve at the top of '
     'the compartment. When product reaches the preset high-level, the OPV automatically closes and '
     'triggers a signal back to the loading skid, causing the flow control valve to close and the '
     'pump to stop — automatically, without any operator action. Overfill is physically impossible '
     'under normal operation.'),
    ('Vapour Recovery System (VRS)',
     'As product enters the bottom of the tanker compartment, displaced vapours are routed back '
     'through a dedicated vapour recovery arm to a vapour recovery unit (VRU) on the skid. The VRU '
     'processes the recovered vapours, typically returning them to the fuel storage tank as liquid. '
     'Recovery efficiency exceeds 95%, effectively eliminating the hazardous Zone 1 atmosphere that '
     'currently exists over the loading area and eliminating the VOC emission issue.'),
    ('Electronic Grounding Verification with Interlock',
     'The loading skid incorporates an earth monitoring system that continuously verifies electrical '
     'continuity of the bond between the loading skid and the tanker body throughout the fill. The '
     'system uses a coded signal — not simple continuity — to prevent bypassing with a jumper wire. '
     'If the ground connection is lost at any point during the fill, product flow stops automatically '
     'within milliseconds. Product cannot flow unless a valid ground is detected.'),
    ('Preset / Pre-set Metered Loading',
     'Modern bottom-load skids incorporate a preset meter — the driver enters (or selects from a '
     'preset) the volume to be delivered to each compartment. The skid controls the fill precisely, '
     'stopping automatically at the specified volume. This provides accurate volume measurement for '
     'invoicing and stock reconciliation.'),
    ('Ground-Level Emergency Shut-off',
     'A clearly marked, accessible emergency stop (E-stop) button is located at each loading bay '
     'at ground level. Activation immediately closes all flow control valves and stops all pumps. '
     'In a top-fill scenario, emergency shut-off often requires the operator to climb down from '
     'height while the spill continues; a bottom-fill E-stop is immediate and at hand.'),
]

for comp_title, comp_body in components:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(comp_title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    add_body(doc, comp_body, space_before=0, space_after=4, size=10)

add_heading(doc, '5.3  Side-by-Side Comparison', level=2, color=MID, size=13, space_before=12)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header cells
h_bad  = tbl.cell(0, 0)
h_good = tbl.cell(0, 1)
set_cell_bg(h_bad,  BG_RED)
set_cell_bg(h_good, BG_GRN)

p = h_bad.paragraphs[0]
r = p.add_run('Current — Top Fill System')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RED

p = h_good.paragraphs[0]
r = p.add_run('Proposed — Bottom Load Skid System')
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREEN

bad_items = [
    'Overfill reliant on operator vigilance — no automatic shut-off',
    'Open hatch during fill — vapours vented to atmosphere continuously',
    'Manual bonding cable — no verification, no interlock, frequently skipped',
    'Operator works at 2.7–3.5 m height every fill, every day',
    'Hose disconnection under flow causes immediate spill',
    'No product metering integrated into loading safety system',
    'Zone 1 hazardous atmosphere maintained throughout fill',
    'NPI VOC reporting likely required',
    'Non-compliant with AS 1940:2017 Cl. 5.5 (grounding interlock)',
]

good_items = [
    'Automatic overfill shut-off via OPV + electronic interlock — overfill impossible',
    'Closed tanker throughout fill — vapours captured by VRS (>95% recovery)',
    'Electronic grounding verification with coded signal and pump interlock',
    'All operations at ground level — fall hazard eliminated',
    'Dry-break API coupler — no spill on disconnect',
    'Preset metered delivery — accurate, auditable volume control',
    'Zone 2 / minimal hazardous atmosphere — dramatically reduced ignition risk',
    'VOC emissions below reporting thresholds in most configurations',
    'Full compliance with AS 1940:2017, WHS Regulations, ADG Code',
]

row = tbl.add_row()
cell_bad  = row.cells[0]
cell_good = row.cells[1]
set_cell_bg(cell_bad,  WHITE)
set_cell_bg(cell_good, WHITE)

for item in bad_items:
    p = cell_bad.add_paragraph(style='List Bullet')
    r = p.add_run(item)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x7b, 0x1c, 0x1c)

for item in good_items:
    p = cell_good.add_paragraph(style='List Bullet')
    r = p.add_run(item)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x0d, 0x4f, 0x28)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 6. REGULATORY COMPLIANCE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '6.  Regulatory Compliance Benefit')
add_body(doc,
    'Installation of the proposed bottom-load system brings the facility into compliance with the '
    'following specific regulatory obligations that are currently unmet or at risk:'
)

add_risk_table(doc,
    ['Regulation / Standard', 'Current Status', 'Post-Upgrade Status'],
    [
        ('AS 1940:2017 Cl. 5.5 — Earthing/bonding interlock for tanker loading', badge('NON-COMPLIANT', RED), badge('COMPLIANT', GREEN)),
        ('WHS Regulations 2017 Reg. 78 — Working at height controls', badge('AT RISK', AMBER), badge('COMPLIANT', GREEN)),
        ('WHS Act 2011 s.19 — Eliminate or minimise foreseeable risks', badge('NON-COMPLIANT', RED), badge('COMPLIANT', GREEN)),
        ('NPI — VOC emission reporting obligations', badge('LIKELY REPORTABLE', AMBER), badge('BELOW THRESHOLD', GREEN)),
        ('State Environmental Legislation — Uncontrolled release', badge('AT RISK', AMBER), badge('CONTROLLED', GREEN)),
        ('ADG Code — Tanker loading at compliant facility', badge('AT RISK', AMBER), badge('COMPLIANT', GREEN)),
    ]
)

add_callout(doc,
    'Insurance Note',
    'Many industrial and public liability insurers are now requiring evidence of AS 1940 compliance '
    'as a condition of cover for fuel handling operations. Facilities found to be non-compliant at '
    'the time of an incident may face claims being declined or coverage voidance. The bottom-load '
    'upgrade removes this exposure.',
    BG_NVY, MID
)

# ════════════════════════════════════════════════════════════════
# 7. BUSINESS CASE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '7.  Business Case & Return on Investment')

add_heading(doc, '7.1  Cost of the Status Quo', level=2, color=MID, size=13, space_before=10)
add_body(doc,
    'The financial exposure of retaining the current system must be understood in full. The following '
    'categories represent quantifiable or estimable cost pools that the current configuration carries:'
)

add_risk_table(doc,
    ['Cost Category', 'Indicative Range', 'Notes'],
    [
        ('WHS regulatory prosecution (serious incident)', '$50,000 – $3,000,000+', 'Maximum penalties under WHS Act for Category 1 offence (reckless endangerment)'),
        ('Environmental prosecution — uncontrolled fuel release', '$25,000 – $1,000,000+', 'State EPA penalties; remediation costs additional and unbounded'),
        ('Tanker fire / explosion — property & business interruption', '$500,000 – $10,000,000+', 'Estimated facility replacement and BI for medium depot; does not include fatality costs'),
        ('Product losses — overfill & vapour', '$5,000 – $40,000 per year', 'Based on medium-volume throughput; vapour losses are ongoing and continuous'),
        ('Operational delays — incident / near-miss response', 'Variable', 'Loading bay downtime, spill response, emergency services attendance'),
        ('Reputational damage', 'Unquantifiable', 'A serious incident at a fuel depot is publicly reported; long-term damage to brand'),
    ]
)

add_heading(doc, '7.2  Benefits of the Upgrade', level=2, color=MID, size=13, space_before=10)
add_bullet(doc, 'All four EXTREME hazard ratings reduced to LOW — full WHS and environmental compliance achieved.', bold_prefix='Risk elimination:  ')
add_bullet(doc, 'Vapour recovery returns product to storage rather than losing it to atmosphere — direct ongoing fuel saving.', bold_prefix='Product recovery:  ')
add_bullet(doc, 'Bottom-fill loading is significantly faster than top-fill for multi-compartment tankers — reduced loading times per tanker.', bold_prefix='Throughput improvement:  ')
add_bullet(doc, 'Eliminating the requirement to climb tankers in all weather conditions improves relationship with fleet operators.', bold_prefix='Driver satisfaction:  ')
add_bullet(doc, 'Demonstrated AS 1940 compliance strengthens the facility\'s insurance position and may reduce premiums.', bold_prefix='Insurance position:  ')
add_bullet(doc, 'Regulators are progressively tightening requirements on VOC emissions and earthing interlocks. Installing a compliant system now avoids forced reactive compliance spending later.', bold_prefix='Future-proofing:  ')

add_heading(doc, '7.3  Capital Cost Estimate', level=2, color=MID, size=13, space_before=10)
add_body(doc,
    'A detailed cost estimate will be provided separately by [Your Company Name] upon completion of '
    'a site survey and final scope confirmation. Indicative cost ranges for a 2–4 bay medium volume '
    'bottom-load conversion at an Australian site are:'
)

add_risk_table(doc,
    ['Scope Item', 'Indicative Range (AUD)'],
    [
        ('Bottom-load skid (per bay) — coupler, meter, valves, controls', '$45,000 – $85,000 per bay'),
        ('Vapour recovery unit (shared, sized for site)', '$30,000 – $80,000'),
        ('Electronic grounding verification system (per bay)', '$4,000 – $10,000 per bay'),
        ('Civil and mechanical installation', '$20,000 – $60,000'),
        ('Electrical and controls', '$15,000 – $40,000'),
        ('Commissioning, testing, and documentation', '$8,000 – $20,000'),
        (('Total indicative range (2–4 bays)', True, NAVY), ('$180,000 – $450,000', True, NAVY)),
    ]
)

add_callout(doc,
    'Context',
    'The indicative total capital cost represents a fraction of the uninsured and insured exposure '
    'of a single serious incident at the facility. It is a one-time capital investment that eliminates '
    'a continuous, ongoing, and growing regulatory and safety liability.',
    BG_NVY, MID
)

# ════════════════════════════════════════════════════════════════
# 8. RECOMMENDATION
# ════════════════════════════════════════════════════════════════
add_heading(doc, '8.  Recommendation & Proposed Next Steps')
add_body(doc,
    'Based on the hazard assessment and regulatory review above, [Your Company Name] strongly recommends '
    '[Customer Name] proceed with the bottom-load skid conversion as a priority safety and compliance '
    'project. The following steps are proposed:'
)

steps = [
    ('Site Survey and Scope Confirmation',
     '[Your Company Name] to attend site to confirm bay dimensions, existing pipework and valve '
     'arrangements, tanker fleet bottom-fill inlet standards (API 1 inch, 4 inch, etc.), electrical '
     'supply availability, and vapour recovery outlet options. Typically 1 day on site.'),
    ('Detailed Design and Quotation',
     'Preparation of a detailed engineering design, equipment schedule, and fixed-price quotation '
     'based on confirmed site conditions and scope. Includes selection of appropriate API coupler '
     'configuration for [Customer Name]\'s tanker fleet.'),
    ('Regulatory Notifications and Approvals',
     'Identification of any development, building, or dangerous goods licence amendments required. '
     '[Your Company Name] can assist with preparation of documentation for relevant authorities.'),
    ('Equipment Procurement and Fabrication',
     'Order placement and skid fabrication. Lead times vary; typically 8–14 weeks for fabricated '
     'skids from Australian suppliers. Factory acceptance testing (FAT) prior to delivery can be arranged.'),
    ('Installation and Commissioning',
     'Mechanical and electrical installation coordinated to minimise downtime — typically completed '
     'bay by bay to maintain loading throughput. Commissioning, calibration, and operator training included.'),
    ('Documentation and Handover',
     'As-built drawings, equipment manuals, maintenance schedules, safety data, and operator training '
     'records provided at handover. [Your Company Name] can also prepare or review updated Safe Work '
     'Method Statements (SWMS) for the new loading procedure.'),
]

for i, (step_title, step_body) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'Step {i}:  ')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    r2 = p.add_run(step_title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = MID
    add_body(doc, step_body, space_before=0, space_after=4, size=10)

# ════════════════════════════════════════════════════════════════
# 9. CONCLUSION
# ════════════════════════════════════════════════════════════════
add_heading(doc, '9.  Conclusion')

add_body(doc,
    'The current top-fill road tanker loading configuration at [Customer Name]\'s facility carries four '
    'simultaneously active, EXTREME-rated safety hazards — overfill, uncontrolled vapour release, '
    'electrostatic discharge ignition, and working at height — with no engineering controls providing '
    'automatic or reliable mitigation. This is not a theoretical risk profile; it is the documented '
    'cause of fatalities, serious injuries, and major fires across the Australian and international '
    'petroleum distribution industry.'
)
add_body(doc,
    'The proposed bottom-load skid system addresses all four hazard categories through engineered '
    'controls, reducing risk ratings from EXTREME to LOW and achieving compliance with the applicable '
    'Australian Standards and WHS legislation. The capital investment required is proportionate and '
    'predictable, and it is demonstrably smaller than the financial exposure of a single serious incident.'
)
add_body(doc,
    '[Customer Name] has the opportunity to make this transition proactively, from a position of choice '
    'and with time to plan properly — rather than reactively, following an incident. [Your Company Name] '
    'is ready to support this process from site survey through to commissioning and handover.'
)

add_callout(doc,
    'Recommended Action',
    'Authorise a site survey and detailed design engagement with [Your Company Name] to establish the '
    'exact scope, cost, and programme for the bottom-load conversion at [Customer Name]\'s depot.',
    BG_GRN, GREEN
)

# ── Footer line ──
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top_el = OxmlElement('w:top')
top_el.set(qn('w:val'), 'single')
top_el.set(qn('w:sz'), '4')
top_el.set(qn('w:space'), '1')
top_el.set(qn('w:color'), f"{MUTED[0]:02X}{MUTED[1]:02X}{MUTED[2]:02X}")
pBdr.append(top_el)
pPr.append(pBdr)
r = p.add_run('BLS-2026-001 Rev A  |  Prepared by [Your Company Name]  |  May 2026  |  Commercial in Confidence')
r.font.size = Pt(8)
r.font.color.rgb = MUTED

# ════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════
output_path = '/home/user/executive-dashboard/Bottom-Load-Fuel-Skid-Safety-Report.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
